from __future__ import annotations

import json
import re
import shutil
import zipfile
from pathlib import Path

from docx import Document
from pypdf import PdfReader


ROOT = Path(__file__).resolve().parents[1]
DOWNLOADS = Path(
    r"C:\Users\veraveen\Downloads\drive-download-20260721T121439Z-1-001"
)
OUT = ROOT / "source_extract"
ASSET_ROOT = ROOT / "assets" / "source_extract"

SOURCES = [
    DOWNLOADS / "codeEvolution-youtube.docx",
    DOWNLOADS / "revisionnotes_1.docx",
    DOWNLOADS / "youtube-code-evolution-advanced-js.docx",
    DOWNLOADS / "nodeUdemy" / "completeNodeDev.docx",
    DOWNLOADS / "nodeUdemy" / "2.1 PDF-Guide-Node-Andrew-Mead-v3.pdf",
]


def slugify(value: str) -> str:
    value = value.lower().strip()
    value = re.sub(r"[^a-z0-9]+", "_", value)
    return value.strip("_") or "source"


def clean_text(value: str) -> str:
    value = value.replace("\xa0", " ")
    value = re.sub(r"[ \t]+", " ", value)
    return value.strip()


def extract_docx_text(path: Path) -> tuple[str, dict]:
    doc = Document(str(path))
    lines: list[str] = [f"# Extract: {path.name}", ""]
    headings: list[str] = []
    para_count = 0
    table_count = 0

    for para in doc.paragraphs:
        text = clean_text(para.text)
        if not text:
            continue
        para_count += 1
        style = para.style.name if para.style else ""
        if "Heading" in style or style in {"Title", "Subtitle"}:
            level = 1
            match = re.search(r"(\d+)$", style)
            if match:
                level = min(int(match.group(1)) + 1, 6)
            headings.append(text)
            lines.append(f"{'#' * level} {text}")
        else:
            lines.append(text)
        lines.append("")

    for idx, table in enumerate(doc.tables, start=1):
        table_count += 1
        lines.append(f"## Table {idx}")
        rows = []
        for row in table.rows:
            cells = [clean_text(cell.text).replace("\n", "<br>") for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if rows:
            width = max(len(row) for row in rows)
            rows = [row + [""] * (width - len(row)) for row in rows]
            lines.append("| " + " | ".join(rows[0]) + " |")
            lines.append("| " + " | ".join(["---"] * width) + " |")
            for row in rows[1:]:
                lines.append("| " + " | ".join(row) + " |")
        lines.append("")

    meta = {
        "name": path.name,
        "type": "docx",
        "paragraphs": para_count,
        "tables": table_count,
        "headings": headings[:200],
    }
    return "\n".join(lines), meta


def extract_docx_images(path: Path, dest: Path) -> list[dict]:
    copied: list[dict] = []
    dest.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(path) as zf:
        for info in zf.infolist():
            if not info.filename.startswith("word/media/"):
                continue
            src_name = Path(info.filename).name
            target = dest / src_name
            with zf.open(info) as src, target.open("wb") as out:
                shutil.copyfileobj(src, out)
            copied.append(
                {
                    "source": path.name,
                    "file": str(target.relative_to(ROOT)).replace("\\", "/"),
                    "bytes": target.stat().st_size,
                }
            )
    return copied


def extract_pdf(path: Path) -> tuple[str, dict]:
    reader = PdfReader(str(path))
    lines = [f"# Extract: {path.name}", ""]
    page_summaries = []
    for index, page in enumerate(reader.pages, start=1):
        text = clean_text(page.extract_text() or "")
        if not text:
            continue
        page_summaries.append({"page": index, "chars": len(text)})
        lines.append(f"## Page {index}")
        lines.append(text)
        lines.append("")
    meta = {
        "name": path.name,
        "type": "pdf",
        "pages": len(reader.pages),
        "pages_with_text": len(page_summaries),
        "page_summaries": page_summaries,
    }
    return "\n".join(lines), meta


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    ASSET_ROOT.mkdir(parents=True, exist_ok=True)

    manifest: dict[str, object] = {"sources": [], "images": []}
    for source in SOURCES:
        if not source.exists():
            manifest["sources"].append({"name": str(source), "missing": True})
            continue

        slug = slugify(source.stem)
        if source.suffix.lower() == ".docx":
            text, meta = extract_docx_text(source)
            images = extract_docx_images(source, ASSET_ROOT / slug)
            meta["images"] = len(images)
            manifest["images"].extend(images)
        elif source.suffix.lower() == ".pdf":
            text, meta = extract_pdf(source)
        else:
            continue

        output_file = OUT / f"{slug}.md"
        output_file.write_text(text, encoding="utf-8")
        meta["extract_file"] = str(output_file.relative_to(ROOT)).replace("\\", "/")
        manifest["sources"].append(meta)

    (OUT / "manifest.json").write_text(
        json.dumps(manifest, indent=2), encoding="utf-8"
    )
    print(json.dumps(manifest, indent=2)[:4000])


if __name__ == "__main__":
    main()
